# ==============================
# RESUME PARSER UTILITY MODULE
# ==============================
import re
import PyPDF2
from nltk.corpus import stopwords
from docx import Document

# ==============================
# SPACY MODEL INITIALIZATION
# ==============================
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except Exception:
    nlp = None
    SPACY_AVAILABLE = False

stop_words = set(stopwords.words('english'))

# ==============================
# TEXT EXTRACTION FUNCTIONS
# ==============================
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    extracted_text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            extracted_text += page_text + "\n"
    return extracted_text

def extract_text_from_docx(file):
    doc = Document(file)
    extracted_text = ""
    for para in doc.paragraphs:
        extracted_text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extracted_text += cell.text + " "
            extracted_text += "\n"
    return extracted_text

def extract_text_from_txt(file):
    return file.read().decode('utf-8', errors='ignore')

def extract_text(file, filename):
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file)
    elif filename_lower.endswith('.txt'):
        return extract_text_from_txt(file)
    else:
        raise ValueError("Unsupported file type. Upload PDF, DOCX or TXT only.")

# ==============================
# TEXT CLEANING FUNCTION
# ==============================
def clean_text(input_text):
    text_lower = input_text.lower()
    text_cleaned = re.sub(r'[^a-zA-Z0-9\s]', ' ', text_lower)
    words = text_cleaned.split()
    filtered_words = [word for word in words if word not in stop_words and len(word) > 2]
    return filtered_words, text_lower

# ==============================
# NAME EXTRACTION FUNCTION
# ==============================
def extract_name(input_text):
    normalized_text = re.sub(r'\s+', ' ', input_text).strip()

    match = re.search(r'\b([A-Z]{2,}(?:\s+[A-Z]{2,}){1,3})\b', normalized_text[:150])
    if match:
        name = match.group(1).strip()
        if 'RESUME' not in name and 'EDUCATION' not in name:
            return name.title()

    words = normalized_text.split()
    role_keywords = ['Data', 'ML', 'AI', 'Engineer', 'Scientist', 'Developer', 'Analyst']
    for i in range(3, 5):
        candidate_words = words[:i]
        if any(w in role_keywords for w in candidate_words):
            continue
        if all(w[0].isupper() and w.isalpha() and len(w) > 2 for w in candidate_words):
            name = " ".join(candidate_words)
            return name

    if SPACY_AVAILABLE and nlp:
        doc = nlp(normalized_text[:500])
        for ent in doc.ents:
            if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                return ent.text.strip()
    return ""

# ==============================
# CONTACT INFO EXTRACTION
# ==============================
def extract_contact_info(input_text):
    info = {
        "name": extract_name(input_text),
        "email": "Not found",
        "phone": "Not found",
        "github": "github.com/username not found",
        "linkedin": "linkedin.com/in/username not found"
    }

    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', input_text)
    if emails:
        info["email"] = emails[0]

    phone_match = re.search(r'(\+?\d{1,3}[\s\-]?)?(\(?\d{3,5}\)?[\s\-]?)?\d{5}[\s\-]?\d{5}', input_text)
    if phone_match:
        info["phone"] = phone_match.group().strip()

    github_url = re.search(r'github\.com/([A-Za-z0-9_.-]+)', input_text, re.IGNORECASE)
    if github_url:
        info["github"] = f"github.com/{github_url.group(1)}"

    linkedin_url = re.search(r'linkedin\.com/in/([A-Za-z0-9_.-]+)', input_text, re.IGNORECASE)
    if linkedin_url:
        info["linkedin"] = f"linkedin.com/in/{linkedin_url.group(1)}"

    return info

# ==============================
# SECTION EXTRACTION BY HEADING - RAW TEXT
# ==============================
def extract_section_by_heading(full_text, heading_keywords):
    text_lower = full_text.lower()
    start_idx = -1
    for keyword in heading_keywords:
        idx = text_lower.find(keyword.lower())
        if idx!= -1:
            start_idx = idx
            break

    if start_idx == -1:
        return f"No {heading_keywords[0]} section found"

    next_headings = ['education', 'experience', 'skills', 'projects', 'certifications', 'certificates', 'professional', 'contact', 'summary', 'objective', 'programming', 'achievements']
    end_idx = len(full_text)

    for heading in next_headings:
        if heading not in [k.lower() for k in heading_keywords]:
            idx = text_lower.find(heading, start_idx + len(heading_keywords[0]))
            if idx!= -1 and idx < end_idx:
                end_idx = idx

    section_text = full_text[start_idx:end_idx].strip()
    lines = section_text.split('\n')
    content_lines = [line.strip() for line in lines[1:] if line.strip()]
    result = '\n'.join(content_lines[:25])
    return result if result.strip() else f"No content under {heading_keywords[0]}"

def extract_projects(input_text):
    return extract_section_by_heading(input_text, ['Projects', 'Project'])

def extract_certifications(input_text):
    return extract_section_by_heading(input_text, ['Certifications', 'Certificates', 'Courses', 'Certificate'])

# ==============================
# SKILL DATABASE - EXPANDED FOR ALL ROLES
# ==============================
SKILLS_DB = {
    "software_engineer": [
        "java", "c++", "c#", "python","html","css", "javascript", "typescript", "go", "rust", "kotlin", "swift",
        "rest api", "graphql", "docker", "git", "github", "linux", "bash", "system design", "dsa",
        "algorithms", "oop", "dbms", "sql", "postgresql", "mysql", "jenkins", "aws", "azure", "gcp"
    ],
    "ai_ml_engineer": [
        "python", "machine learning", "ml", "deep learning", "dl", "tensorflow", "pytorch", "keras",
        "scikit-learn", "numpy", "pandas", "nlp", "computer vision", "cv", "llm", "large language model",
        "langchain", "langgraph", "langsmith", "openai", "chatgpt", "huggingface", "transformers",
        "bert", "gpt", "rag", "vector database", "pinecone", "weaviate", "faiss", "prompt engineering",
        "fine-tuning", "agentic ai", "reinforcement learning", "mlops", "weights and biases", "wandb"
    ],
    "backend_developer": [
        "python", "django", "flask", "fastapi", "nodejs", "express", "java", "spring boot", "golang",
        "rest api", "graphql", "postgresql", "mysql", "mongodb", "redis", "rabbitmq", "kafka",
        "docker", "kubernetes", "aws", "gcp", "azure", "nginx", "gunicorn", "git"
    ],
    "frontend_developer": [
        "html", "css", "javascript", "typescript", "react", "nextjs", "vue", "angular", "redux",
        "tailwind", "bootstrap", "sass", "webpack", "vite", "responsive design", "rest api"
    ],
    "data_science": [
        "python", "r", "sql", "machine learning", "statistics", "pandas", "numpy", "matplotlib",
        "seaborn", "scikit-learn", "xgboost", "lightgbm", "tableau", "powerbi", "spark", "hadoop",
        "airflow", "data visualization", "eda", "regression", "classification"
    ],
    "devops": [
        "docker", "kubernetes", "k8s", "aws", "azure", "gcp", "terraform", "ansible", "jenkins",
        "github actions", "gitlab ci", "ci/cd", "linux", "bash", "monitoring", "prometheus", "grafana"
    ],
    "cloud": [
        "aws", "amazon web services", "azure", "gcp", "google cloud platform", "ec2", "s3", "lambda",
        "cloudformation", "vpc", "iam", "bigquery", "vertexai", "cloud storage", "compute engine"
    ]
}

ROLE_MAP = {
    "software_engineer": "Software Engineer",
    "ai_ml_engineer": "AI/ML Engineer",
    "backend_developer": "Backend Developer",
    "frontend_developer": "Frontend Developer",
    "data_science": "Data Scientist",
    "devops": "DevOps Engineer",
    "cloud": "Cloud Engineer"
}

# ==============================
# ROLE-WISE MINIMUM REQUIRED SKILLS - MAIN ONLY
# ==============================
ROLE_MIN_REQUIRED_SKILLS = {
    "Software Engineer": ["python", "java", "javascript", "git", "sql", "dsa", "oops", "linux"],
    "AI/ML Engineer": ["python", "machine learning", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch"],
    "Backend Developer": ["python", "django", "flask", "sql", "rest api", "git"],
    "Frontend Developer": ["html", "css", "javascript", "react", "git"],
    "Data Scientist": ["python", "sql", "pandas", "numpy", "machine learning", "statistics"],
    "DevOps Engineer": ["docker", "kubernetes", "aws", "linux", "git", "ci/cd"],
    "Cloud Engineer": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform"]
}

# ==============================
# SKILL EXTRACTION
# ==============================
def extract_skills(words):
    combined_text = " ".join(words).lower()
    found_skills = set()
    for category, skills in SKILLS_DB.items():
        for skill in skills:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, combined_text, re.IGNORECASE):
                found_skills.add(skill)
    return sorted(list(found_skills))

# ==============================
# ROLE PREDICTION - UPDATED
# ==============================
def predict_role(skills):
    skill_set = {s.lower() for s in skills}
    category_scores = {}

    for category, skill_list in SKILLS_DB.items():
        count = sum(1 for skill in skill_list if skill.lower() in skill_set)
        if category == "ai_ml_engineer":
            count *= 2.0
        if category == "software_engineer":
            count *= 1.5
        if category == "cloud":
            count *= 1.3
        category_scores[category] = count

    if max(category_scores.values()) == 0:
        return "Fresher"
    return ROLE_MAP[max(category_scores, key=category_scores.get)]

# ==============================
# ATS SCORE CALCULATION
# ==============================
def calculate_ats(input_text, skills):
    score = 0
    text_lower = input_text.lower()

    sections = ['education', 'experience', 'project', 'skill', 'certification']
    found_sections = sum(1 for s in sections if s in text_lower)
    score += (found_sections / len(sections)) * 20

    total_db_skills = len(set(sum(SKILLS_DB.values(), [])))
    skill_ratio = len(skills) / total_db_skills if total_db_skills > 0 else 0
    score += min(skill_ratio * 25, 25)

    impact_keywords = ['%', 'increased', 'improved', 'reduced', 'accuracy', 'deployed', 'built', 'optimized', 'scaled']
    impact_count = sum(1 for kw in impact_keywords if kw in text_lower)
    score += min(impact_count * 1.5, 15)

    projects = extract_projects(input_text)
    if "No Projects section found" not in projects:
        score += 15

    certs = extract_certifications(input_text)
    if "No Certifications section found" not in certs:
        score += 10

    if '•' in input_text or '-' in input_text:
        score += 10

    role_keywords = ['software', 'developer', 'engineer', 'coding', 'algorithm', 'architecture']
    keyword_count = sum(1 for kw in role_keywords if kw in text_lower)
    score += min(keyword_count * 0.8, 5)

    return int(min(score, 100))


# ==============================
# MISSING SKILLS SUGGESTION
# ==============================
def missing_skills(skills, predicted_role):
    user_skills = {s.lower() for s in skills}
    min_required_list = ROLE_MIN_REQUIRED_SKILLS.get(predicted_role, [])

    # Count how many minimum required skills user has
    found_min_skills = [s for s in min_required_list if s.lower() in user_skills]

    # Logic: If user has 4+ minimum skills, show motivational quote instead of empty
    min_threshold = 4

    if len(found_min_skills) >= min_threshold:
        # Quote based on role
        quotes = {
            "Software Engineer": "✨ All core skills present! You're job-ready 💪",
            "AI/ML Engineer": "🚀 Core ML skills strong! Ready to build AI products",
            "Backend Developer": "⚡ Backend basics covered! Time to scale systems",
            "Frontend Developer": "🎨 Frontend fundamentals strong! Build amazing UIs",
            "Data Scientist": "📊 Core data skills ready! Analyze and conquer",
            "DevOps Engineer": "🔧 DevOps basics covered! Deploy with confidence",
            "Cloud Engineer": "☁️ Cloud fundamentals strong! Scale to the sky",
            "Fresher": "🌱 Great start! Keep learning and building"
        }
        return [quotes.get(predicted_role, "✨ All essential skills present!")]

    # If less than 4, show only missing minimum skills
    missing = [skill.title() for skill in min_required_list if skill.lower() not in user_skills]
    return missing[:6]

